"""
============================================================================
 Fix BiDi (Bidirectional) Text in Word Documents.
============================================================================
 Reads a .docx file with mixed Hebrew (RTL) and English (LTR) text,
 fixes the directionality so paragraphs flow correctly, and saves to
 a new .docx file.

 Usage:
     python fix_bidi_docx.py <input.docx> <output.docx>
============================================================================
"""

import re
import sys
import copy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# -----------------------------------------------------------------------
#  Unicode helpers
# -----------------------------------------------------------------------

_HEB_RANGE = re.compile(r'[\u0590-\u05FF]')
_ENG_RANGE = re.compile(r'[A-Za-z]')


def _has_hebrew(text: str) -> bool:
    return bool(_HEB_RANGE.search(text))


def _has_english(text: str) -> bool:
    return bool(_ENG_RANGE.search(text))


def _is_hebrew_char(ch: str) -> bool:
    return bool(_HEB_RANGE.match(ch))


def _is_english_char(ch: str) -> bool:
    return bool(_ENG_RANGE.match(ch))


def _is_neutral(ch: str) -> bool:
    return not _is_hebrew_char(ch) and not _is_english_char(ch)


# -----------------------------------------------------------------------
#  Segment mixed text into (text, direction) chunks
# -----------------------------------------------------------------------

def _segment_bidi(text: str) -> list[tuple[str, str]]:
    """
    ====================================================================
     Split text into segments of consecutive Hebrew or English chars.
     Neutral chars (digits, punctuation, spaces) attach to the
     preceding segment's direction, defaulting to RTL.
    ====================================================================
    """
    if not text:
        return []

    segments: list[tuple[str, str]] = []
    current = []
    current_dir = 'rtl'

    for ch in text:
        if _is_hebrew_char(ch):
            if current_dir != 'rtl' and current:
                segments.append((''.join(current), current_dir))
                current = []
            current_dir = 'rtl'
            current.append(ch)
        elif _is_english_char(ch):
            if current_dir != 'ltr' and current:
                segments.append((''.join(current), current_dir))
                current = []
            current_dir = 'ltr'
            current.append(ch)
        else:
            # Neutral — attach to current direction
            current.append(ch)

    if current:
        segments.append((''.join(current), current_dir))

    return segments


# -----------------------------------------------------------------------
#  XML helpers for paragraph and run BiDi properties
# -----------------------------------------------------------------------

def _set_paragraph_rtl(paragraph) -> None:
    """
    ====================================================================
     Set paragraph direction to RTL and alignment to right.
    ====================================================================
    """
    pPr = paragraph._p.get_or_add_pPr()

    # Set BiDi on paragraph
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
    bidi.set(qn('w:val'), '1')

    # Right-to-left alignment
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pPr.append(jc)
    jc.set(qn('w:val'), 'right')


def _set_run_rtl(run) -> None:
    """
    ====================================================================
     Mark a run as RTL.
    ====================================================================
    """
    rPr = run._r.get_or_add_rPr()
    rtl = rPr.find(qn('w:rtl'))
    if rtl is None:
        rtl = OxmlElement('w:rtl')
        rPr.append(rtl)
    rtl.set(qn('w:val'), '1')


def _set_run_ltr(run) -> None:
    """
    ====================================================================
     Mark a run as LTR and remove any RTL flag.
    ====================================================================
    """
    rPr = run._r.get_or_add_rPr()
    rtl = rPr.find(qn('w:rtl'))
    if rtl is not None:
        rPr.remove(rtl)


def _copy_run_format(source_run, target_run) -> None:
    """
    ====================================================================
     Copy font formatting from source run to target run.
    ====================================================================
    """
    src_rPr = source_run._r.find(qn('w:rPr'))
    if src_rPr is not None:
        new_rPr = copy.deepcopy(src_rPr)
        old_rPr = target_run._r.find(qn('w:rPr'))
        if old_rPr is not None:
            target_run._r.remove(old_rPr)
        target_run._r.insert(0, new_rPr)


# -----------------------------------------------------------------------
#  Core fix logic
# -----------------------------------------------------------------------

def _fix_paragraph(paragraph) -> None:
    """
    ====================================================================
     Fix BiDi for a single paragraph:
     1. Set paragraph direction to RTL (if it contains Hebrew).
     2. Re-segment runs so each run has a uniform direction.
     3. Mark each run with the correct RTL/LTR property.
    ====================================================================
    """
    full_text = paragraph.text
    if not full_text or not full_text.strip():
        return

    has_heb = _has_hebrew(full_text)
    if not has_heb:
        return

    # Set paragraph-level RTL
    _set_paragraph_rtl(paragraph)

    # Collect all run texts and a reference run for formatting
    ref_run = paragraph.runs[0] if paragraph.runs else None
    if ref_run is None:
        return

    # Segment the full paragraph text
    segments = _segment_bidi(full_text)
    if not segments:
        return

    # Remove existing runs from the XML
    p_element = paragraph._p
    for run in list(paragraph.runs):
        p_element.remove(run._r)

    # Create new runs per segment
    for seg_text, seg_dir in segments:
        new_run = paragraph.add_run(seg_text)
        _copy_run_format(ref_run, new_run)
        if seg_dir == 'rtl':
            _set_run_rtl(new_run)
        else:
            _set_run_ltr(new_run)


def fix_bidi_docx(input_path: str, output_path: str) -> None:
    """
    ====================================================================
     Fix BiDi in a Word document and save to output path.
    ====================================================================
    """
    doc = Document(input_path)

    for paragraph in doc.paragraphs:
        _fix_paragraph(paragraph)

    # Also fix paragraphs inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _fix_paragraph(paragraph)

    doc.save(output_path)
    print(f'Fixed: {input_path} -> {output_path}')


# -----------------------------------------------------------------------
#  CLI
# -----------------------------------------------------------------------

def main() -> None:
    if len(sys.argv) != 3:
        print('Usage: python fix_bidi_docx.py <input.docx> <output.docx>')
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    if not Path(input_path).exists():
        print(f'Error: File not found: {input_path}')
        sys.exit(1)

    fix_bidi_docx(input_path=input_path, output_path=output_path)


if __name__ == '__main__':
    main()
